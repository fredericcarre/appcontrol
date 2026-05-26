#!/usr/bin/env python3
"""
Build the Invivoo-branded Word reference template used by Pandoc to
style the generated strategy.docx.

Run once (or whenever you want to refresh the template):
    python3 docs/build-reference-docx.py

Output:
    docs/invivoo-reference.docx

How it is applied:
    docs/build-docx.sh detects this file and runs:
        pandoc strategy.md -o strategy.docx --reference-doc=invivoo-reference.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ---------- Palette (mirrors the HTML doc) --------------------------------
INK         = RGBColor(0x0F, 0x17, 0x2A)   # near-black navy
INK_SOFT    = RGBColor(0x33, 0x41, 0x55)   # body text
MUTED       = RGBColor(0x64, 0x74, 0x8B)
LINE        = RGBColor(0xE2, 0xE8, 0xF0)
ACCENT      = RGBColor(0x0F, 0x76, 0x6E)   # teal (primary brand)
ACCENT_DEEP = RGBColor(0x13, 0x4E, 0x4A)
ACCENT_SOFT = RGBColor(0xCC, 0xFB, 0xF1)
WARN        = RGBColor(0xB4, 0x53, 0x09)
WARN_SOFT   = RGBColor(0xFE, 0xF3, 0xC7)
OK          = RGBColor(0x16, 0x65, 0x34)
OK_SOFT     = RGBColor(0xDC, 0xFC, 0xE7)
DANGER      = RGBColor(0xB9, 0x1C, 0x1C)
DANGER_SOFT = RGBColor(0xFE, 0xE2, 0xE2)
INDIGO      = RGBColor(0x43, 0x38, 0xCA)
INDIGO_SOFT = RGBColor(0xE0, 0xE7, 0xFF)


# ---------- XML helpers ---------------------------------------------------
def _set_shading(paragraph, hex_color: str) -> None:
    """Apply a solid background fill to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_left_border(paragraph, hex_color: str, size: int = 24) -> None:
    """Add a thick coloured left border to a paragraph (Pull-quote style)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))   # 1/8 pt units
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_bottom_border(paragraph, hex_color: str, size: int = 6) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)


def _rgb_hex(rgb: RGBColor) -> str:
    return "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])


def _ensure_style(doc, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    styles = doc.styles
    if name in [s.name for s in styles]:
        return styles[name]
    return styles.add_style(name, style_type)


# ---------- Build the template -------------------------------------------
def build(out_path: Path) -> None:
    doc = Document()

    # ---- Page setup -----------------------------------------------------
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(0.9)

    # ---- Default body font ---------------------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK_SOFT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    # ---- Title ----------------------------------------------------------
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(34)
    title.font.bold = True
    title.font.color.rgb = ACCENT_DEEP
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    # ---- Subtitle -------------------------------------------------------
    if "Subtitle" in [s.name for s in doc.styles]:
        sub = doc.styles["Subtitle"]
        sub.font.name = "Calibri"
        sub.font.size = Pt(15)
        sub.font.italic = False
        sub.font.color.rgb = MUTED
        sub.paragraph_format.space_after = Pt(18)

    # ---- Heading 1 (chapter) -------------------------------------------
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT_DEEP
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    # ---- Heading 2 -----------------------------------------------------
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = ACCENT
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ---- Heading 3 -----------------------------------------------------
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12.5)
    h3.font.bold = True
    h3.font.color.rgb = INK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # ---- Quote (used by Pandoc for > blockquotes) ----------------------
    quote = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else _ensure_style(doc, "Quote")
    quote.font.name = "Calibri"
    quote.font.size = Pt(11.5)
    quote.font.italic = False
    quote.font.color.rgb = ACCENT_DEEP
    quote.paragraph_format.left_indent = Cm(0.3)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(8)

    # ---- Custom paragraph styles ---------------------------------------
    def _add_callout(name: str, fg: RGBColor, bg: RGBColor, border: RGBColor):
        style = _ensure_style(doc, name)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = fg
        pf = style.paragraph_format
        pf.left_indent = Cm(0.3)
        pf.right_indent = Cm(0.3)
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)
        return style, fg, bg, border

    _add_callout("CalloutOK",      OK,     OK_SOFT,     OK)
    _add_callout("CalloutWarn",    WARN,   WARN_SOFT,   WARN)
    _add_callout("CalloutDanger",  DANGER, DANGER_SOFT, DANGER)
    _add_callout("CalloutIndigo",  INDIGO, INDIGO_SOFT, INDIGO)

    # Pull quote — sarcelle bar on left, light teal bg
    _add_callout("PullQuote",      ACCENT_DEEP, ACCENT_SOFT, ACCENT)
    _add_callout("PullQuoteWarn",  WARN, WARN_SOFT, WARN)

    # Draft banner — red box across full width
    draft = _ensure_style(doc, "DraftBanner")
    draft.base_style = doc.styles["Normal"]
    draft.font.name = "Calibri"
    draft.font.size = Pt(10)
    draft.font.bold = True
    draft.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    draft.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    draft.paragraph_format.space_before = Pt(2)
    draft.paragraph_format.space_after = Pt(14)

    # Eyebrow / kicker (small uppercase muted)
    eye = _ensure_style(doc, "Eyebrow")
    eye.base_style = doc.styles["Normal"]
    eye.font.name = "Calibri"
    eye.font.size = Pt(9)
    eye.font.bold = True
    eye.font.color.rgb = ACCENT
    eye.font.all_caps = True
    eye.paragraph_format.space_after = Pt(0)

    # Timeline phase banner (teal/indigo)
    phase_a = _ensure_style(doc, "TimelinePhaseA")
    phase_a.base_style = doc.styles["Heading 2"]
    phase_a.font.color.rgb = ACCENT
    phase_a.font.size = Pt(13)
    phase_a.font.all_caps = True
    phase_a.paragraph_format.space_before = Pt(14)

    phase_b = _ensure_style(doc, "TimelinePhaseB")
    phase_b.base_style = doc.styles["Heading 2"]
    phase_b.font.color.rgb = INDIGO
    phase_b.font.size = Pt(13)
    phase_b.font.all_caps = True
    phase_b.paragraph_format.space_before = Pt(14)

    # ---- Header (Invivoo + Confidentiel) -------------------------------
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hdr_p = header.paragraphs[0]
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hdr_p.add_run("INVIVOO  ·  Confidentiel")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = ACCENT_DEEP

    tab_stops = hdr_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.6))
    hdr_p.add_run("\t")
    r2 = hdr_p.add_run("Répondre avec sérénité aux exigences de DORA")
    r2.font.name = "Calibri"
    r2.font.size = Pt(9)
    r2.font.color.rgb = MUTED
    _set_bottom_border(hdr_p, _rgb_hex(LINE), size=6)

    # ---- Footer (copyright + draft + page number) ---------------------
    footer = section.footer
    foot_p = footer.paragraphs[0]
    foot_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = foot_p.add_run("© 2026 Invivoo — Tous droits réservés")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = MUTED

    foot_p.paragraph_format.tab_stops.add_tab_stop(Cm(8.3), WD_ALIGN_PARAGRAPH.CENTER)
    foot_p.paragraph_format.tab_stops.add_tab_stop(Cm(16.6), WD_ALIGN_PARAGRAPH.RIGHT)

    foot_p.add_run("\t")
    draft_run = foot_p.add_run("DRAFT v0.9 · Ne pas diffuser")
    draft_run.font.name = "Calibri"
    draft_run.font.size = Pt(8.5)
    draft_run.font.bold = True
    draft_run.font.color.rgb = DANGER

    foot_p.add_run("\t")

    # Page number field
    page_run = foot_p.add_run()
    page_run.font.name = "Calibri"
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    page_run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    page_run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_end)
    page_run.add_text(" / ")
    pages_run = foot_p.add_run()
    pages_run.font.name = "Calibri"
    pages_run.font.size = Pt(8.5)
    pages_run.font.color.rgb = MUTED
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    pages_run._r.append(fld_begin2)
    instr2 = OxmlElement("w:instrText")
    instr2.text = "NUMPAGES"
    pages_run._r.append(instr2)
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    pages_run._r.append(fld_end2)

    _set_bottom_border(foot_p, _rgb_hex(LINE), size=0)  # placeholder
    pPr = foot_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), _rgb_hex(LINE))
    pBdr.append(top)
    pPr.append(pBdr)

    # ---- First-page header/footer (cover): minimal ---------------------
    first_header = section.first_page_header
    first_header.paragraphs[0].add_run("")  # blank — cover is self-branded
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].add_run("")

    # ---- Save ----------------------------------------------------------
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")


if __name__ == "__main__":
    here = Path(__file__).resolve().parent
    build(here / "invivoo-reference.docx")
